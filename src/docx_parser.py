# src/docx_parser.py
"""Module trích xuất dữ liệu từ file DOCX - Phương án sử dụng vốn"""

import re
from typing import Dict, Any, Optional
from docx import Document
from src.utils import clean_text, parse_number


class DocxParser:
    """Class để parse file DOCX phương án sử dụng vốn"""
    
    def __init__(self, file_path: str):
        """
        Khởi tạo parser
        
        Args:
            file_path: Đường dẫn file DOCX
        """
        self.doc = Document(file_path)
        self.paragraphs = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        self.text_content = "\n".join(self.paragraphs)
        self.tables = self._extract_tables()
    
    def _extract_tables(self) -> list:
        """Trích xuất tất cả bảng từ document"""
        tables_data = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        return tables_data
    
    def _extract_value_from_line(self, line: str, separator: str = ':') -> str:
        """Trích xuất giá trị sau dấu phân cách"""
        if separator in line:
            parts = line.split(separator, 1)
            if len(parts) > 1:
                return parts[1].strip()
        return ""
    
    def _find_paragraph_containing(self, keywords: list) -> Optional[str]:
        """Tìm đoạn văn chứa từ khóa"""
        for keyword in keywords:
            for para in self.paragraphs:
                if keyword.lower() in para.lower():
                    return para
        return None
    
    def _extract_number_from_text(self, text: str) -> float:
        """Trích xuất số từ text"""
        if not text:
            return 0.0
        
        # Tìm số đầu tiên (có thể có dấu phân cách)
        # Ví dụ: "5.000.000.000" hoặc "5,000,000,000"
        match = re.search(r'(\d[\d.,]*)', text)
        if match:
            number_str = match.group(1)
            return parse_number(number_str)
        return 0.0
    
    def extract_customer_info(self) -> Dict[str, str]:
        """Trích xuất thông tin khách hàng"""
        name = ""
        cccd = ""
        address = ""
        phone = ""
        
        # Tìm thông tin từ phần I. Thông tin chung về khách hàng
        found_first_person = False
        
        for i, para in enumerate(self.paragraphs):
            # Bỏ qua nếu chưa đến phần thông tin khách hàng
            if "I. Thông tin chung" in para:
                continue
            
            # Tìm thông tin người đầu tiên (chỉ lấy 1 lần)
            if not found_first_person:
                # Tìm họ tên từ dòng "1. Họ và tên:"
                if para.startswith("1. Họ và tên:") or "1. Họ và tên:" in para:
                    # Extract name từ pattern "1. Họ và tên: Tên -Sinh năm:"
                    match = re.search(r'1\.\s*Họ và tên:\s*([^-\n]+)', para)
                    if match:
                        name = match.group(1).strip()
                        found_first_person = True
            
            # Chỉ lấy CCCD của người đầu tiên
            if found_first_person and not cccd:
                if "CMND/CCCD" in para:
                    # Extract CCCD từ nhiều patterns
                    match = re.search(r'(?:CMND/CCCD|CCCD)[^:]*:\s*(\d{9,12})', para)
                    if match:
                        cccd = match.group(1).strip()
            
            # Lấy địa chỉ đầu tiên
            if found_first_person and not address and "Nơi cư trú:" in para:
                address = self._extract_value_from_line(para, "Nơi cư trú:")
            
            # Lấy số điện thoại đầu tiên
            if found_first_person and not phone and "Số điện thoại:" in para:
                match = re.search(r'Số điện thoại:\s*(\d{10,11})', para)
                if match:
                    phone = match.group(1).strip()
            
            # Dừng khi gặp mục 2 hoặc phần II
            if found_first_person and (para.startswith("2. Họ và tên:") or "II." in para):
                break
        
        return {
            'name': name or 'Khách hàng',
            'cccd': cccd,
            'address': address,
            'phone': phone
        }
    
    def extract_loan_info(self) -> Dict[str, Any]:
        """Trích xuất thông tin khoản vay"""
        total_need = 0.0
        equity = 0.0
        loan_amount = 0.0
        interest_rate = 8.5
        loan_term = 60
        purpose = "Mua nhà"
        
        for para in self.paragraphs:
            # Tổng nhu cầu vốn - lấy số đầy đủ
            if para.startswith("1. Tổng nhu cầu vốn:") or "1. Tổng nhu cầu vốn:" in para:
                # Extract full number
                match = re.search(r'Tổng nhu cầu vốn:\s*([\d.,]+)', para)
                if match:
                    total_need = parse_number(match.group(1))
            
            # Vốn đối ứng - lấy từ dòng có "Vốn đối ứng tham gia"
            if "Vốn đối ứng tham gia" in para:
                match = re.search(r':\s*([\d.,]+)', para)
                if match:
                    equity = parse_number(match.group(1))
            
            # Vốn vay - lấy từ dòng "Vốn vay Agribank số tiền:"
            if "Vốn vay Agribank số tiền:" in para:
                match = re.search(r'số tiền:\s*([\d.,]+)', para)
                if match:
                    loan_amount = parse_number(match.group(1))
            
            # Mục đích vay
            if para.startswith("Mục đích vay:"):
                purpose = self._extract_value_from_line(para, "Mục đích vay:")
                if not purpose:
                    purpose = "Mua nhà"
            
            # Thời hạn vay - lấy chính xác
            if para.startswith("Thời hạn vay:"):
                match = re.search(r'Thời hạn vay:\s*(\d+)\s*tháng', para)
                if match:
                    loan_term = int(match.group(1))
            
            # Lãi suất - lấy chính xác
            if "Lãi suất:" in para and "Thời hạn vay:" in para:
                match = re.search(r'Lãi suất:\s*(\d+[.,]?\d*)\s*%', para)
                if match:
                    interest_rate = float(match.group(1).replace(',', '.'))
        
        # Tính equity ratio
        equity_ratio = (equity / total_need * 100) if total_need > 0 else 0
        
        return {
            'purpose': purpose,
            'total_need': total_need,
            'equity': equity,
            'loan_amount': loan_amount,
            'equity_ratio': equity_ratio,
            'interest_rate': interest_rate,
            'loan_term': loan_term,
            'payment_frequency': 'Tháng'
        }
    
    def extract_collateral_info(self) -> Dict[str, Any]:
        """Trích xuất thông tin tài sản bảo đảm"""
        asset_type = "Bất động sản"
        market_value = 0.0
        asset_address = ""
        ltv = 70.0
        legal_docs = ""
        
        # Tìm trong phần 5. Tài sản bảo đảm
        in_collateral_section = False
        
        for para in self.paragraphs:
            if "5. Tài sản bảo đảm" in para or "Tài sản bảo đảm:" in para:
                in_collateral_section = True
                continue
            
            if in_collateral_section:
                # Tài sản 1
                if "Tài sản 1:" in para:
                    # Extract type và value
                    match = re.search(r'Tài sản 1:\s*([^\.]+).*Giá trị:\s*([\d.,]+)', para)
                    if match:
                        asset_type = match.group(1).strip()
                        market_value = parse_number(match.group(2))
                    else:
                        # Thử pattern khác
                        if "Giá trị:" in para:
                            market_value = self._extract_number_from_text(para.split("Giá trị:")[1])
                
                # Địa chỉ tài sản
                if "Địa chỉ:" in para and in_collateral_section:
                    asset_address = self._extract_value_from_line(para, "Địa chỉ:")
                
                # LTV
                if "LTV" in para or "Tỷ lệ cho vay" in para:
                    match = re.search(r'(\d+[.,]?\d*)\s*%', para)
                    if match:
                        ltv = float(match.group(1).replace(',', '.'))
                
                # Giấy chứng nhận
                if "Giấy chứng nhận" in para:
                    legal_docs = para
                
                # Dừng khi gặp phần III
                if "III." in para:
                    break
        
        return {
            'asset_type': asset_type,
            'market_value': market_value,
            'asset_address': asset_address,
            'ltv': ltv,
            'legal_docs': legal_docs if legal_docs else 'Sổ đỏ/Giấy chứng nhận'
        }
    
    def extract_financial_info(self) -> Dict[str, float]:
        """Trích xuất thông tin tài chính"""
        monthly_income = 0.0
        monthly_expense = 0.0
        other_debt = 0.0
        
        # Tìm trong phần 4. Nguồn trả nợ
        for para in self.paragraphs:
            # Tổng thu nhập
            if "Tổng thu nhập" in para and "tháng" in para.lower():
                monthly_income = self._extract_number_from_text(para)
            
            # Thu nhập từ lương
            if "Thu nhập từ lương:" in para:
                monthly_income = self._extract_number_from_text(para)
            
            # Tổng chi phí
            if "Tổng chi phí hàng tháng:" in para:
                monthly_expense = self._extract_number_from_text(para)
            
            # Chi phí sinh hoạt
            if "Chi phí sinh hoạt" in para and monthly_expense == 0:
                monthly_expense = self._extract_number_from_text(para)
        
        return {
            'monthly_income': monthly_income,
            'monthly_expense': monthly_expense,
            'other_debt': other_debt
        }
    
    def parse_full_document(self) -> Dict[str, Any]:
        """
        Parse toàn bộ document
        
        Returns:
            Dictionary chứa tất cả thông tin đã trích xuất
        """
        customer_info = self.extract_customer_info()
        loan_info = self.extract_loan_info()
        collateral_info = self.extract_collateral_info()
        financial_info = self.extract_financial_info()
        
        return {
            'customer_info': customer_info,
            'loan_info': loan_info,
            'collateral_info': collateral_info,
            'financial_info': financial_info,
            'raw_text': self.text_content
        }

